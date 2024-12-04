import streamlit as st
from docx import Document
import os
from pathlib import Path
import base64
import io
from docx.shared import Inches
import pandas as pd
from PIL import Image
import numpy as np

# キャッシュ機能を使用して処理を効率化
@st.cache_data
def get_document_preview(doc):
    """ドキュメントのプレビューテキストを取得"""
    preview_text = []
    for para in doc.paragraphs[:5]:  # 最初の5段落をプレビュー
        if para.text.strip():
            preview_text.append(para.text)
    return "\n".join(preview_text)

@st.cache_data
def get_document_images(doc):
    """ドキュメント内の画像を取得"""
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            images.append(image_data)
    return images

def merge_word_documents(files, keep_styles=False, keep_images=False, progress_bar=None):
    """
    アップロードされたWord文書を1つのファイルにまとめる
    """
    try:
        merged_document = Document()
        total_files = len(files)
        
        for idx, uploaded_file in enumerate(files):
            # プログレスバーの更新
            if progress_bar is not None:
                progress = (idx + 1) / total_files
                progress_bar.progress(progress)
            
            # ファイルの内容を読み込む
            doc = Document(io.BytesIO(uploaded_file.read()))
            
            # ファイル名をセクション見出しとして追加
            merged_document.add_heading(uploaded_file.name, level=1)
            
            if keep_styles:
                # スタイルを保持してコピー
                for element in doc.element.body:
                    merged_document.element.body.append(element)
            else:
                # 通常のテキストとしてコピー
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        merged_document.add_paragraph(paragraph.text)
            
            if keep_images:
                # 画像を保持してコピー
                images = get_document_images(doc)
                for img_data in images:
                    img_stream = io.BytesIO(img_data)
                    merged_document.add_picture(img_stream, width=Inches(6))
            
            # セクション区切りを追加
            merged_document.add_paragraph('=' * 50)
            merged_document.add_paragraph()
        
        doc_stream = io.BytesIO()
        merged_document.save(doc_stream)
        return doc_stream.getvalue()
    
    except Exception as e:
        st.error(f"ファイルの結合中にエラーが発生しました: {str(e)}")
        return None

def get_download_link(binary_doc, filename):
    """ダウンロードリンクを生成する"""
    try:
        b64 = base64.b64encode(binary_doc).decode()
        return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="download-button">ダウンロード</a>'
    except Exception as e:
        st.error(f"ダウンロードリンクの生成中にエラーが発生しました: {str(e)}")
        return None

def main():
    try:
        # ページ設定
        st.set_page_config(
            page_title="Word文書結合アプリ",
            page_icon="📄",
            layout="wide",
            initial_sidebar_state="expanded"
        )
        
        # アプリタイトルと説明
        st.title("📄 Word文書結合アプリ")
        st.write("複数のWord文書(.docx)を1つのファイルにまとめることができます。")
        
        # サイドバーのオプション設定
        with st.sidebar:
            st.header("⚙️ オプション設定")
            keep_styles = st.checkbox("スタイルを保持", value=True, 
                                    help="元のドキュメントのフォント、色、サイズなどのスタイルを保持します")
            keep_images = st.checkbox("画像を保持", value=True,
                                    help="元のドキュメント内の画像を保持します")
            
            st.markdown("---")
            st.markdown("### 💡 Tips")
            st.markdown("""
            - ドラッグ＆ドロップでファイルをアップロード
            - 順序は自由に変更可能
            - プレビューで内容を確認してから結合
            """)
        
        # メインコンテンツ
        uploaded_files = st.file_uploader(
            "結合したいWordファイル(.docx)を選択してください",
            type=["docx"],
            accept_multiple_files=True
        )
        
        if uploaded_files:
            st.write(f"📁 アップロードされたファイル数: {len(uploaded_files)}")
            
            # ファイル順序の管理
            st.subheader("📋 ファイルの順序を変更")
            file_order_df = pd.DataFrame({
                'ファイル名': [f.name for f in uploaded_files],
                '順序': list(range(1, len(uploaded_files) + 1))
            })
            
            edited_df = st.data_editor(
                file_order_df,
                hide_index=True,
                use_container_width=True
            )
            
            # ファイルの順序を更新
            ordered_files = []
            for _, row in edited_df.sort_values('順序').iterrows():
                for f in uploaded_files:
                    if f.name == row['ファイル名']:
                        ordered_files.append(f)
            
            # プレビュー表示
            st.subheader("👀 ファイルプレビュー")
            selected_file = st.selectbox(
                "プレビューするファイルを選択",
                options=[f.name for f in ordered_files]
            )
            
            for f in ordered_files:
                if f.name == selected_file:
                    try:
                        doc = Document(io.BytesIO(f.read()))
                        f.seek(0)  # ファイルポインタをリセット
                        
                        with st.expander(f"プレビュー: {f.name}"):
                            preview_text = get_document_preview(doc)
                            st.text(preview_text)
                            
                            if keep_images:
                                images = get_document_images(doc)
                                if images:
                                    st.write("📸 含まれる画像:")
                                    cols = st.columns(3)
                                    for idx, img_data in enumerate(images[:3]):
                                        with cols[idx]:
                                            img = Image.open(io.BytesIO(img_data))
                                            st.image(img, use_column_width=True)
                    except Exception as e:
                        st.error(f"プレビューの生成中にエラーが発生しました: {str(e)}")
            
            # 出力ファイル名の入力
            output_filename = st.text_input(
                "出力ファイル名を入力してください",
                value="merged_document.docx"
            )
            
            # 結合ボタン
            if st.button("🔄 ファイルを結合", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                try:
                    status_text.text("🔄 ファイルを結合中...")
                    merged_doc = merge_word_documents(
                        ordered_files,
                        keep_styles=keep_styles,
                        keep_images=keep_images,
                        progress_bar=progress_bar
                    )
                    
                    if merged_doc:
                        status_text.text("✅ 完了！")
                        st.markdown(
                            get_download_link(merged_doc, output_filename),
                            unsafe_allow_html=True
                        )
                        st.success("ファイルの結合が完了しました！")
                    
                except Exception as e:
                    st.error(f"エラーが発生しました: {str(e)}")
                
                finally:
                    progress_bar.empty()
        
        # 使い方の説明
        with st.expander("❓ 使い方"):
            st.write("""
            1. サイドバーでオプションを設定します
               - スタイルを保持: 元のドキュメントのフォーマットを維持
               - 画像を保持: ドキュメント内の画像を保持
            
            2. 「ファイルをアップロード」から結合したいWord文書を選択
            
            3. ファイルの順序を変更する場合は、表の「順序」列を編集
            
            4. プレビューで内容を確認
            
            5. 出力ファイル名を設定
            
            6. 「ファイルを結合」ボタンをクリック
            
            注意事項:
            - アップロードできるファイルは.docx形式のみです
            - 大きなファイルの場合、処理に時間がかかる場合があります
            - スタイルを保持すると、一部のフォーマットが崩れる可能性があります
            """)

        # CSSでダウンロードボタンをスタイリング
        st.markdown("""
            <style>
            .download-button {
                display: inline-block;
                padding: 0.5em 1em;
                background-color: #4CAF50;
                color: white;
                text-decoration: none;
                border-radius: 4px;
                transition: background-color 0.3s;
            }
            .download-button:hover {
                background-color: #45a049;
            }
            </style>
        """, unsafe_allow_html=True)

    except Exception as e:
        st.error(f"アプリケーションの実行中にエラーが発生しました: {str(e)}")

if __name__ == "__main__":
    main()